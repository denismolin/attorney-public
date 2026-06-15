"""Extraction du texte des pièces jointes.

Jalon 1 : PDF natif (PyMuPDF), DOCX, RTF, TXT.
Jalon 2 : PDF scanné → OCR via Claude Vision (parsing/ocr.py) si ANTHROPIC_API_KEY est configurée.
"""
from __future__ import annotations

from pathlib import Path
from typing import Callable

from config import DATA_DIR
from parsing.eml_parser import Attachment

# Seuil : en dessous de N caractères extraits d'un PDF, on considère que c'est
# un scan (texte non sélectionnable) → needs_ocr.
_PDF_MIN_CHARS = 40


def _abs_path(att: Attachment) -> Path:
    """stored_path est relatif au parent de data/ ; on reconstruit l'absolu."""
    return (DATA_DIR.parent / att.stored_path).resolve()


def _extract_pdf(path: Path) -> tuple[str, bool]:
    """Retourne (texte, needs_ocr). PyMuPDF lit le texte natif."""
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    text = "\n".join(parts).strip()
    needs_ocr = len(text) < _PDF_MIN_CHARS
    return ("" if needs_ocr else text), needs_ocr


def _extract_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    paras = [p.text for p in document.paragraphs]
    # Tableaux DOCX éventuels (le détail tabulaire fin viendra au jalon 2).
    for table in document.tables:
        for row in table.rows:
            paras.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paras).strip()


def _extract_rtf(path: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw).strip()


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def extract_text(
    att: Attachment,
    api_key: str | None = None,
    log: Callable[[str], None] | None = None,
) -> tuple[str, bool]:
    """Extrait le texte d'une PJ. Retourne (texte, needs_ocr).

    Si api_key est fournie et que le PDF est scanné, tente l'OCR via Claude Vision.
    Si l'OCR réussit, needs_ocr est posé à False (le texte est indexable).
    Les types non pris en charge renvoient ("", False) sans erreur.
    """
    parts = extract_text_parts(att, api_key=api_key, log=log)
    if not parts:
        return "", False
    first_text, needs_ocr = parts[0]
    # Concat all parts for callers that don't support multi-part (legacy path).
    if len(parts) > 1:
        combined = "\n\n".join(t for t, _ in parts if t)
        return combined, needs_ocr
    return first_text, needs_ocr


def extract_text_parts(
    att: Attachment,
    api_key: str | None = None,
    log: Callable[[str], None] | None = None,
) -> list[tuple[str, bool]]:
    """Extrait le texte d'une PJ. Retourne une liste de (texte, needs_ocr).

    Pour les PDF scannés de plus de 30 pages, retourne une entrée par tranche de
    30 pages. Pour tous les autres cas, retourne une liste à un seul élément.
    Les types non pris en charge retournent [("", False)].
    """
    path = _abs_path(att)
    if not path.exists():
        return [("", False)]

    suffix = path.suffix.lower()
    mime = att.mime.lower()

    try:
        if suffix == ".pdf" or mime == "application/pdf":
            text, needs_ocr = _extract_pdf(path)
            if needs_ocr and api_key:
                from parsing.ocr import extract_text_vision_parts
                ocr_parts = extract_text_vision_parts(path, api_key, log=log)
                valid = [(t, False) for t in ocr_parts if len(t) >= _PDF_MIN_CHARS]
                if valid:
                    return valid
            return [(text, needs_ocr)]
        if suffix == ".docx" or "wordprocessingml" in mime:
            return [(_extract_docx(path), False)]
        if suffix == ".rtf" or mime in ("application/rtf", "text/rtf"):
            return [(_extract_rtf(path), False)]
        if suffix in (".txt", ".md") or mime.startswith("text/"):
            return [(_extract_txt(path), False)]
    except Exception as exc:  # pragma: no cover - robustesse indexation
        print(f"[attachments] extraction échouée {path.name}: {exc}")
        return [("", False)]

    return [("", False)]
